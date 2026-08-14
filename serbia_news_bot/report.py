"""报告生成（Word .docx）：使馆内政简报体 + 来源 + 原文链接。"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from .ai import AIVerdict
from .article import ParsedArticle
from .config import REPORTS_DIR

_WEEKDAYS = "一二三四五六日"
_META_COLOR = RGBColor(0x55, 0x55, 0x55)
_LINK_COLOR = RGBColor(0x05, 0x63, 0xC1)


@dataclass
class ReportItem:
    article: ParsedArticle
    verdict: AIVerdict


def _date_zh(target_date: date) -> str:
    weekday = _WEEKDAYS[target_date.weekday()]
    return f"{target_date.year}年{target_date.month}月{target_date.day}日（星期{weekday}）"


def _set_run_font(run, *, size_pt: float, color: RGBColor | None = None, bold: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _add_meta_line(doc: Document, label: str, value: str, *, link: bool = False) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(2)
    label_run = para.add_run(f"{label}：")
    _set_run_font(label_run, size_pt=10.5, color=_META_COLOR)
    value_run = para.add_run(value)
    _set_run_font(value_run, size_pt=10.5, color=_LINK_COLOR if link else _META_COLOR)
    if link:
        value_run.underline = True


def build_docx(target_date: date, items: list[ReportItem], scanned: int) -> Document:
    doc = Document()

    title = doc.add_heading("塞尔维亚在野党动态每日专报", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(_date_zh(target_date))
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        _set_run_font(sub.runs[0], size_pt=12, color=RGBColor(0x66, 0x66, 0x66))

    doc.add_heading("一、内政", level=1)

    if not items:
        doc.add_paragraph(f"今日（{_date_zh(target_date)}）未收录符合口径的在野党相关硬新闻。")
        return doc

    for idx, item in enumerate(items, start=1):
        ver = item.verdict
        heading = ver.title_zh.strip() or item.article.title
        doc.add_heading(heading, level=2)

        for para_text in ver.summary_zh.strip().split("\n"):
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                if p.runs:
                    _set_run_font(p.runs[0], size_pt=12)

        _add_meta_line(doc, "来源", item.article.source_name)
        _add_meta_line(doc, "链接", item.article.url, link=True)

        if idx < len(items):
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)

    # scanned 仅用于日志侧统计，报告正文不再展示元数据
    _ = scanned
    return doc


def write_report(target_date: date, items: list[ReportItem], scanned: int) -> Path:
    path = Path(REPORTS_DIR)
    path.mkdir(parents=True, exist_ok=True)
    out = path / f"report_{target_date.isoformat()}.docx"
    doc = build_docx(target_date, items, scanned)
    doc.save(out)
    return out
